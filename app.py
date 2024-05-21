import pyttsx3
import speech_recognition as sr
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from flask import Flask, request, jsonify
from flask import render_template
import requests
import time
import openai
import os
from openai import OpenAI
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import pandas as pd  #?
import re
import requests
import time
import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.pdfgen import canvas
import unicodedata
import re
import PyPDF2
from docx import Document
requests.packages.urllib3.disable_warnings()

engine = pyttsx3.init()
recognizer = sr.Recognizer()
#engine.setProperty('volume', 1.0)
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[1].id)  # 0 for male voice, 1 for female voice
engine.setProperty('rate', 150) 
app = Flask(__name__)

# resume = "Komal Bhole B-Tech IT, VIIT Pune I'm excited to start my career and bring my skills onto the table. I'm a quick learner and eager to take on new challenges. I'm looking forward to connecting with other professionals in the industry and exploring new opportunities. komal.22010509@viit.ac.in 7499203665 Pune, India 02 June, 2003 linkedin.com/in/komal-bhole-4904a2204 github.com/KOMAL616 EDUCATION B-Tech - Information Technology Vishwakarma Institute of Information Technology 01/2021 - Present, 9.19 HSC Board Brijilal Biyani Science College, Amravati SKILLSC C++ Python(Numpy, Pandas, Matplotlib) Java HTML Git/GitHub CSS DBMS(MySql) Operating System Figma Data Structure and Algorithms Oops Php Computer Networking PERSONAL PROJECTS Vegetable Ordering Website (06/2022 - 12/2022) Enable customers to order vegetables online Convenient and hassle-free way to purchase vegies. Technologies used: Html, Css, Javascript, Php, MySql. Health App (Healthify (01/2023 - 04/2023) Booking appointments from any hybrid location.. Unique QR code consisting of all the minute details of patient AI based chatbot for clearing basic doubts. Technologies Used : Java , Xml, Kotlin, Firebase. Built a React CI/CD Pipeline using CircleCI 1. Collaborate with the team to develop and implement machine learning algorithms and models 2. Assist in collecting, cleaning, and analyzing large datasets to derive actionable insights 3. Develop and maintain data preprocessing pipelines for efficient data handling 4. Contribute to the design and development of NLP models for text classification and sentiment analysis 5. Research cutting-edge ML techniques and stay updated with the latest advancements in the field 6. Assist in building and improving ML models for personalized recommendations and content filtering 7. Help with the deployment and monitoring of ML models in production environments Join us and make a significant impact in revolutionizing the way people access knowledge and grow personally through our innovative platform. Apply now and take the first step towards an exciting career in Machine Learning with MentorBoxx! Skill(s) requiredData Analytics Machine Learning Natural Language Processing (NLP)Python"
# Position = "Software Developer"
# job_description = "Design, develop, and test software solutions Collaborate with cross-functional teams to define, design, and ship new features Strong programming skills in languages like Java, Python, or C++"
# Company_Name = "Google"


questions = ""
conversation_log = []
n = 0
conversation_user = ""
timestamps = []
next_question = ""
system_prompt = ""
# system_prompt = [  {"role": "system", "content": f"Hi GPT, you are  acting as an interviewer in this exercise. You are to interview me for {Position} at {Company_Name} with the job description posted as {job_description}. My resume is as follows:{resume}In This exercise you are  conducting my interview and current conversation log is as {conversation_log} if it is empty means it is first question,   Whenever, I would be conversing with you, the ChatGPT model and not as the interviewer,  After the interview, you are to assess me based on the following five aspects:Technical Skills: Assessing the candidate's proficiency in the specific technical skills required for the job. Behavioural Skills: Evaluating communication, teamwork, adaptability, and other interpersonal skills. Problem-Solving and Critical Thinking: Assessing analytical thinking, decision-making, and the ability to handle challenges. Cultural Fit: Examining alignment with the company's values, culture, and the ability to work within a team. Work Experience and Achievements: Reviewing past experiences, achievements, and contributions to understand the candidate's capabilities and potential impact. Hence, please note to ask questions such that all of these aspects can be evaluated and considered in the question set.  currently you have asked {n}  number of questions and the number of questions  should strictly not exceed 3, Once the questions exceed the limit. Simply respond with the following response: EXIT_0 at the start of the responce and greet the candidate and announce that interview is over and instruct him to wait for feedack" }]


client = OpenAI(api_key='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')
@app.route("/")
def home():
    return render_template("index.html")


import datetime
@app.route("/start", methods=['POST'])
def start_interview():
    # def upload_file(file):
    #     if file:
    #         if file.filename.endswith('.pdf'):
    #             reader = PyPDF2.PdfFileReader(file)
    #             text = ''
    #             for page in range(reader.getNumPages()):
    #                 text += reader.getPage(page).extractText()
    #             return text
    #         elif file.filename.endswith('.docx'):
    #             doc = Document(file)
    #             text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
    #             return text

    #     return 'No file uploaded'

    def upload_file(file):
        if file and file.filename:
            if file.filename.endswith('.pdf'):
                reader = PyPDF2.PdfReader(file)
                text = ''
                for page in range(len(reader.pages)):
                    text += reader.pages[page].extract_text()
                return text
            elif file.filename.endswith('.docx'):
                doc = Document(file)
                text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
                return text
        return 'No file uploaded'

    name = request.form.get('name')
    job_title = request.form.get('job_title')
    job_description = request.form.get('job_description')
    company_name = request.form.get('company_name')
    interviewer = request.form.get('interviewer')
    resume = upload_file(request.files.get('resume'))
    print(resume)

    global questions, n, conversation_log, system_prompt, timestamps, next_question
    if interviewer.lower() == 'lalit':
        instrutions =   """
        Technical Skills (40%): Focus primarily on evaluating the candidate's technical proficiency. Ask detailed technical questions related to the job requirements, and provide scenarios to assess practical application.
        Behavioural Skills (20%): Include some questions to gauge basic behavioural competencies, such as communication and teamwork, but with less emphasis compared to technical skills.
        Problem-Solving and Critical Thinking (20%): Incorporate problem-solving scenarios related to technical challenges faced in the role. Assess the candidate's ability to analyze problems and propose effective solutions.
        Cultural Fit (10%): Ask a few questions about the candidate's approach to teamwork and collaboration, but with less depth compared to technical and problem-solving aspects.
        Work Experience and Achievements (10%): Briefly discuss the candidate's past experiences to understand their technical contributions and achievements.
        """

    elif interviewer.lower() == 'komal':
        instructions = """ Technical Skills (15%): Ask a few basic technical questions to ensure the candidate meets minimum requirements, but with less emphasis compared to other aspects.ehavioural Skills (30%): Place significant emphasis on evaluating communication, teamwork, and adaptability through situational questions and role-playing scenarios.Problem-Solving and Critical Thinking (15%): Include some problem-solving scenarios to assess the candidate's ability to handle challenges, but with less complexity compared to other interviewers.Cultural Fit (30%): Focus heavily on discussing the candidate's alignment with company values, culture, and their ability to contribute positively to the team dynamic.Work Experience and Achievements (10%): Briefly discuss the candidate's past experiences to understand their overall fit within the company culture and team environment."""
    elif interviewer.lower() == 'mayank':
        instructions = """ candidate meets the basic requirements, but with less emphasis compared to other aspects.
        Behavioural Skills (15%): Ask a few questions related to communication and teamwork, but with minimal depth compared to problem-solving and critical thinking.
        Problem-Solving and Critical Thinking (40%): Place a heavy emphasis on presenting complex scenarios and assessing the candidate's ability to analyze problems, think critically, and propose innovative solutions.
        Cultural Fit (15%): Discuss the candidate's alignment with company culture briefly, focusing on their approach to handling challenges within a team environment.
        Work Experience and Achievements (10%): Briefly review the candidate's past experiences to understand how they have applied problem-solving skills in previous roles.
        """

    elif interviewer.lower() == 'abhay':
        instructions = """ Technical Skills (30%): Assess the candidate's technical capabilities with a mix of theoretical and practical questions.
        Behavioural Skills (25%): Place significant emphasis on evaluating communication, adaptability, and teamwork skills through scenario-based questions.
        Problem-Solving and Critical Thinking (20%): Include a moderate level of problem-solving scenarios to assess analytical thinking and decision-making abilities.
        Cultural Fit (15%): Probe the candidate's alignment with company values and culture through situational questions related to team dynamics and work environment.
        Work Experience and Achievements (10%): Briefly discuss the candidate's past experiences to understand their overall impact and contributions.
        """

    initial_prompt = interview_script = f"""
    Hi ChatGPT, you are to act as an interviewer in this exercise and conduct my interviw as per the instructions {instructions}.  My resume is as follows:

    {resume}
    
    This exercise would consist of you conducting my interview consisting of questions in the range of 30 - 50. The number of questions should strictly not exceed this. Whenever, I would be conversing with you, the ChatGPT model and not as the interviewer, I would prompt you as: 

    system: 

    Other prompts during the interviews which would be answers to your questions, would be given as prompts to you without any such prefix.

    After the interview, you are to assess me based on the following five aspects:


    Technical Skills: Assessing the candidate's proficiency in the specific technical skills required for the job.
    Behavioural Skills: Evaluating communication, teamwork, adaptability, and other interpersonal skills.
    Problem-Solving and Critical Thinking: Assessing analytical thinking, decision-making, and the ability to handle challenges.
    Cultural Fit: Examining alignment with the company's values, culture, and the ability to work within a team.
    Work Experience and Achievements: Reviewing past experiences, achievements, and contributions to understand the candidate's capabilities and potential impact.


    Hence, please note to ask questions such that all of these aspects can be evaluated and considered in the question set. Once you are done with the interview questions. Simply respond with the following response:
    EXIT 0
    """
    conversation_log.append({'prompt': initial_prompt})



    system_prompt = [  {"role": "system", "content": f"Hi ChatGPT, here’s an interview log conducted by you for the {job_description} pertaining to {job_title} as: {conversation_log}. Now, I need you to continue the interview as per the initial interview instructions. You are strictly instructed to ask only one question at a time regardless the complexity and depth. The number of questions asked until now is {n}. Hence, ask a follow up question. Strictly make sure not to repeat any question and conduct the entire interview as per instructions in the prompts given earlier to you in the conversation log. If you feel that you want to conclude the interview as per the number of questions or feel that all the assessment aspects have been covered up, please reply with just the response “EXIT 0” and nothing else." }]
    

    question = generate_question(system_prompt)
    print(question)
    conversation_log.append({"role": "assistant", "content": question})
    #print(conversation_log)
    speak(question)



    continue_interview = True
    while continue_interview:

        if n>3:
            last_question = "Hello there! I must say, your interview went  well. Now, please patiently wait for the feedback. It won't take long, and I'm sure you'll find it insightful. Great job again, and thank you for your time during the interview!"
            print(last_question)
            speak(last_question)
            feedback = generate_feedback(conversation_log)
            parsed_feedback = parse_feedback(feedback)
            generate_pdf(parsed_feedback)
            continue_interview = False
        else:
            n = n+1
            start_time = time.time()  
            timestamps.append(start_time)
            response = listen()
            print(response)
            end_time = time.time()  # Record the end time
            timestamps.append(end_time)
            conversation_log.append({"role": "user", "content": response})
            
            next_question = generate_question(system_prompt)
            print(next_question)
            if n<4:
                speak(next_question)

            conversation_log.append({"role": "assistant", "content": next_question})
    return '', 200

def generate_question(prompt):
    chat_completion = client.chat.completions.create(
          messages = prompt,
          model="gpt-3.5-turbo",
          max_tokens=4096,
          top_p=0.9,)
    
    question = chat_completion.choices[0].message.content
    return question

def listen():
    with sr.Microphone() as source:
        print("Listening...")
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source)
        try:
            text = recognizer.recognize_google(audio)
            return text
        except sr.UnknownValueError:
            return "Sorry, I did not understand that."
        except sr.RequestError as e:
            return f"Sorry, an error occurred. {e}"

def count_grammar_errors(conversation):
    total_errors = 0
    language_tool_url = "https://languagetool.org/api/v2/check"
    data = {
        "text": conversation,
        "language": "en-US"
    }
    response = requests.post(language_tool_url, data=data, verify=False)
    if response.ok:
        matches = response.json()["matches"]
        total_errors = len(matches)
    return total_errors

def calculate_fluency(conversation_duration, error_count):
    max_errors_per_minute = 10
    conversation_duration_minutes = conversation_duration / 60
    max_allowable_errors = max_errors_per_minute * conversation_duration_minutes
    fluency_percentage = 100 - ((error_count / max_allowable_errors) * 100)
    return max(0, fluency_percentage)

def speak(text):
    engine.say(text)
    engine.startLoop(False)
    engine.iterate()
    engine.endLoop()

def generate_feedback(conversation):
    prompt = f"""
    To provide structured and comprehensive feedback for a candidate's interview based on the following conversation log: {conversation},

    Evaluate the candidate across these five key aspects :

    **1. Technical Skills**:
    **Score**: [Enter score out of 100]
    - **Sr. No**: 1 - **Topic Name**: Example Skill - **Remarks**: Strengths; Needs improvement - **Resources**: Example resource, Another resource
    ...
    **Summary**: Provide a summary of the candidate's performance in technical skills, highlighting both strengths and weaknesses.

    **2. Behavioural Skills**:
    **Score**: [Enter score out of 100]
    - **Sr. No**: 1 - **Topic Name**: Example Skill - **Remarks**: Strengths; Needs improvement - **Resources**: Example resource, Another resource
    ...
    **Summary**: Provide a summary of the candidate's behavioural skills, focusing on areas of both proficiency and required development.

    **3. Problem-Solving and Critical Thinking**:
    **Score**: [Enter score out of 100] 
    - **Sr. No**: 1 - **Topic Name**: Example Skill - **Remarks**: Strengths; Needs improvement  - **Resources**: Example resource, Another resource
    ...
    **Summary**: Discuss the candidate’s problem-solving skills and provide insights into their decision-making capabilities.

    **4. Cultural Fit**:
    **Score**: [Enter score out of 100]
    - **Sr. No**: 1 - **Topic Name**: Example Skill - **Remarks**: Strengths; Needs improvement - **Resources**: Example resource, Another resource
    ...
    **Summary**: Summarize how well the candidate fits into the company culture and team environment.

    **5. Work Experience and Achievements**:
    **Score**: [Enter score out of 100]
    - **Sr. No**: 1 - **Topic Name**: Example Skill - **Remarks**: Strengths; Needs improvement - **Resources**: Example resource, Another resource
    ...
    **Summary**: Analyze the candidate’s previous experiences and achievements, emphasizing the impact and breadth of their work.

    strictly provide feedback in above mentioned format This structured feedback will be visualized in a PDF format, where each section corresponds to a separate table on a dedicated page.
    for your reference Evaluation Metrics Table   would be a table consisting of the following columns: Sr. No ,Topic Name, Remarks and Resources 
    Topic Name refers to the topic you evaluated in the section Remarks is Your remarks in bullet points for the topic; try to include both positive and negative remarks here . Keep this as descriptive as possible but short. You are free to use phrases to mind the word limit.  
    Resources is a list of websites, books or any helpful resource that can be used by the candidate to improve on his weak aspects and points for the topic evaluated. provide actual references to existing resources for that skill provide each resource in th form of hyperlink 
    Summary should be a summary of a candidate's performance for the section mentioning the shortcomings, positive aspects, wrong responses and overall justification for the grading.
    """
    system = [{"role": "system", "content": prompt}]
    feedback = generate_question(system)
    return feedback

def generate_pdf(feedback):
    filename = "Candidate_Feedback.pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter)
    content = []
    styles = getSampleStyleSheet()
    for section in feedback:
        # Section Title
        section_title = Paragraph(f"<b>{section['Section Name']}</b>", styles['Title'])
        content.append(section_title)

        # Section Score
        score = Paragraph(f"Score: {section['Score']}/100", styles['BodyText'])
        content.append(score)
        
        # Table for Evaluation Metrics
        table_data = [['Sr. No', 'Topic Name', 'Remarks', 'Resources']]
        
        for item in section['Evaluation Metrics Table']:
            remarks = Paragraph('<br/>'.join(item['Remarks']), styles['BodyText'])
            
            # Update links to be clickable
            resources_links = []
            for resource in item['Resources']:
                # Check if the line contains a hyperlink format
                if '[' in resource and ']' in resource and '(' in resource and ')' in resource:
                    # Extract the display text and URL
                    start_link = resource.index('[') + 1
                    end_link = resource.index(']')
                    start_url = resource.index('(') + 1
                    end_url = resource.index(')')
                    
                    link_text = resource[start_link:end_link]
                    url = resource[start_url:end_url]
                    
                    # Form an HTML-like link with an anchor tag
                    link_html = f'<link href="{url}">{link_text}</link>'
                    resources_links.append(link_html)
                else:
                    resources_links.append(resource)
            
            resources_paragraph = Paragraph('<br/>'.join(resources_links), styles['BodyText'])
            table_data.append([item['Sr. No'], item['Topic Name'], remarks, resources_paragraph])
        
        table = Table(table_data, splitByRow=True)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.gray),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (1, 1), (-1, -1), colors.lightgrey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]))

        table.hAlign = 'LEFT'
        content.append(table)

        # Summary Paragraph
        summary = Paragraph(section['Summary'], styles['BodyText'])
        content.append(summary)
        content.append(Spacer(1, 12))  # Adding space after each section
        content.append(PageBreak())  # Page break after each section

    doc.build(content)
# Example feedback structured list

def parse_feedback(feedback_string):
    # Define the structured data format
    feedback_data = []

    # Split feedback into sections based on known headings
    # Use more robust split regex that looks for number followed by period and any text up to colon or newline
    sections = re.split(r'\n(?=\*\*(\d+\.\s[^*]+?)\*\*:)', feedback_string)
    sections = [s.strip() for s in sections if s.strip()]

    # Process each section
    for section in sections:
        # Initialize default values
        
        section_name = "Undefined Section"  # Default if no match is found
        score = 0

        # Extracting the section name
        section_name_match = re.match(r'\*\*(\d+\.\s*[^*]+?)\*\*:', section)
        if section_name_match:
            section_name = section_name_match.group(1)

        # Extracting the score
        score_match = re.search(r'\*\*Score\*\*:\s*(\d+)', section)
        if score_match:
            score = score_match.group(1)

        # Extract points for the Evaluation Metrics Table
        table_content = []
        points = re.findall(r'- \*\*Sr\. No\*\*:\s*(\d+)\s*- \*\*Topic Name\*\*:\s*([^*]+)\s*- \*\*Remarks\*\*:\s*([^*]+)\s*- \*\*Resources\*\*:\s*([^\n]+)', section)
        for point in points:
            sr_no, topic_name, remarks, resources = point
            remarks_list = remarks.split('; ')
            resources_list = [resource.strip() for resource in resources.split(',')]
            table_content.append({'Sr. No': sr_no, 'Topic Name': topic_name, 'Remarks': remarks_list, 'Resources': resources_list})

        # Extract the summary
        summary = ""  # Default summary if no match is found
        summary_match = re.search(r'\*\*Summary\*\*:\s*([^*]+)', section)
        if summary_match:
            summary = summary_match.group(1)

        # Append structured data for current section
        if section_name != 'Undefined Section':
            feedback_data.append({
                'Section Name': section_name,
                'Score': score,
                'Evaluation Metrics Table': table_content,
                'Summary': summary
            })
    
    return feedback_data
if __name__ == '_main_':
    app.run(debug=True)